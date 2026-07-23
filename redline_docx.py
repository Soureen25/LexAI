"""
redline_docx.py
───────────────
Takes extracted document text + the AI's risk list and produces a redlined
.docx where every flagged clause is shown struck-through in red with the
suggested replacement inserted immediately after in blue underline — exactly
like a lawyer's markup.

Clauses that can't be reliably located (fuzzy-match score below threshold)
are collected in an appendix at the end so nothing is silently dropped.

Usage:
    from redline_docx import build_redline_docx

    docx_bytes = build_redline_docx(document_text, risks_list, doc_name="NDA.pdf")
    # returns raw .docx bytes ready for st.download_button
"""

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from rapidfuzz import fuzz, process

# ── Thresholds ────────────────────────────────────────────────────────────────
FUZZY_MATCH_THRESHOLD = 72   # minimum score (0-100) to auto-locate a clause
MIN_CLAUSE_CHARS      = 15   # ignore trivially short clauses
MAX_CLAUSE_CHARS      = 600  # truncate very long clauses before matching

# ── Colours ───────────────────────────────────────────────────────────────────
RED   = RGBColor(0x99, 0x1B, 0x1B)
BLUE  = RGBColor(0x1E, 0x3A, 0x8A)
NAVY  = RGBColor(0x0F, 0x21, 0x37)
GREY  = RGBColor(0x6B, 0x72, 0x80)
AMBER = RGBColor(0x92, 0x40, 0x0E)
GREEN = RGBColor(0x06, 0x5F, 0x46)

RISK_COLORS = {"HIGH": RED, "MEDIUM": AMBER, "LOW": GREEN}


# ── XML helpers ───────────────────────────────────────────────────────────────
def _strike(run):
    """Apply strikethrough formatting to a run."""
    rpr = run._r.get_or_add_rPr()
    strike = OxmlElement("w:strike")
    strike.set(qn("w:val"), "true")
    rpr.append(strike)
    return run


def _add_run(para, text, bold=False, italic=False, color=None,
             underline=False, strike=False, size_pt=10.5):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    if color:
        run.font.color.rgb = color
    if strike:
        _strike(run)
    run.font.size = Pt(size_pt)
    return run


def _h_rule(doc):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E7EB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── Clause matching ───────────────────────────────────────────────────────────
def _split_paragraphs(text: str) -> list[tuple[str, int]]:
    """
    Split document text into (paragraph_text, start_char_offset) tuples,
    skipping blank lines. Used as the corpus for fuzzy matching.
    """
    paras = []
    pos = 0
    for line in text.splitlines():
        stripped = line.strip()
        if stripped:
            paras.append((stripped, pos))
        pos += len(line) + 1   # +1 for the newline character
    return paras


def _find_best_match(clause: str, paragraphs: list[tuple[str, int]]):
    """
    Fuzzy-match a risk clause against all document paragraphs.
    Returns (best_paragraph_text, score) or (None, 0) if nothing is good enough.
    """
    clause = clause.strip()[:MAX_CLAUSE_CHARS]
    if len(clause) < MIN_CLAUSE_CHARS:
        return None, 0

    corpus = [p[0] for p in paragraphs]
    result = process.extractOne(
        clause, corpus,
        scorer=fuzz.partial_ratio,
        score_cutoff=FUZZY_MATCH_THRESHOLD,
    )
    if not result:
        # retry with token_set_ratio for longer/reordered clauses
        result = process.extractOne(
            clause, corpus,
            scorer=fuzz.token_set_ratio,
            score_cutoff=FUZZY_MATCH_THRESHOLD,
        )
    if result:
        return result[0], result[1]
    return None, 0


# ── Core builder ─────────────────────────────────────────────────────────────
def build_redline_docx(document_text: str, risks: list, doc_name: str = "document") -> bytes:
    """
    Build a redlined .docx from flat extracted document text + AI risk list.

    Since the input is extracted text (not a native .docx), we reconstruct
    the document paragraph-by-paragraph and inject redline markup wherever a
    flagged clause is matched. This is transparent by design — the user can
    see exactly what changed and why, and accept/reject each edit in Word.

    Args:
        document_text : raw text extracted from the original file
        risks         : list of dicts with keys clause / risk_level /
                        why_its_a_risk / suggestion
        doc_name      : original filename, used in the cover page

    Returns:
        Raw .docx bytes.
    """
    doc      = Document()
    section  = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── Cover page ────────────────────────────────────────────────────────────
    title = doc.add_heading("Redlined Legal Document", level=1)
    title.runs[0].font.color.rgb = NAVY

    meta_para = doc.add_paragraph()
    _add_run(meta_para, f"Source document: ", bold=True, color=GREY, size_pt=9)
    _add_run(meta_para, doc_name, color=GREY, size_pt=9)
    meta_para.add_run("\n")
    _add_run(meta_para, f"Generated: ", bold=True, color=GREY, size_pt=9)
    _add_run(meta_para, datetime.now().strftime("%d %b %Y, %H:%M"), color=GREY, size_pt=9)
    meta_para.add_run("\n")
    _add_run(meta_para, f"Risks applied: ", bold=True, color=GREY, size_pt=9)
    _add_run(meta_para, str(len(risks)), color=GREY, size_pt=9)

    legend = doc.add_paragraph()
    _add_run(legend, "  ORIGINAL CLAUSE (deleted)  ", color=RED, strike=True, size_pt=9)
    _add_run(legend, "   ", size_pt=9)
    _add_run(legend, "  SUGGESTED REPLACEMENT  ", color=BLUE, underline=True, size_pt=9)

    disclaimer = doc.add_paragraph()
    _add_run(
        disclaimer,
        "⚠  This document is AI-generated. All suggested changes should be reviewed "
        "by a qualified legal professional before execution.",
        italic=True, color=GREY, size_pt=8.5,
    )

    _h_rule(doc)
    doc.add_paragraph()

    # ── Build clause lookup map: matched_paragraph_text → [risk, ...] ─────────
    paragraphs = _split_paragraphs(document_text)
    ORDER      = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    sorted_risks = sorted(risks, key=lambda r: ORDER.get(str(r.get("risk_level","")).upper(), 3))

    matched_map   = {}   # para_text -> list of risks that matched it
    unmatched     = []   # risks we couldn't locate

    for risk in sorted_risks:
        clause = str(risk.get("clause", "")).strip()
        best_para, score = _find_best_match(clause, paragraphs)
        if best_para:
            matched_map.setdefault(best_para, []).append((risk, score))
        else:
            unmatched.append(risk)

    # ── Reconstruct document with redline injections ──────────────────────────
    body_heading = doc.add_heading("Document (with redlines applied)", level=2)
    body_heading.runs[0].font.color.rgb = NAVY

    for para_text, _offset in paragraphs:
        if para_text in matched_map:
            risks_for_para = matched_map[para_text]

            for risk, score in risks_for_para:
                level      = str(risk.get("risk_level", "LOW")).upper()
                why        = risk.get("why_its_a_risk", "")
                suggestion = risk.get("suggestion", "")
                level_color = RISK_COLORS.get(level, GREY)

                # Risk label line
                label_para = doc.add_paragraph()
                _add_run(label_para, f"⚑  {level} RISK  ", bold=True,
                          color=level_color, size_pt=8.5)
                _add_run(label_para, f"(match confidence: {score:.0f}%)  · ",
                          color=GREY, size_pt=8)
                _add_run(label_para, why, italic=True, color=GREY, size_pt=8.5)

                # Redline paragraph: struck-through original + blue suggestion
                redline_para = doc.add_paragraph()
                _add_run(redline_para, para_text + "  ",
                          color=RED, strike=True, size_pt=10.5)
                _add_run(redline_para, suggestion,
                          color=BLUE, underline=True, bold=True, size_pt=10.5)
        else:
            # Normal paragraph — no changes
            p = doc.add_paragraph(para_text)
            p.runs[0].font.size = Pt(10.5) if p.runs else None

    # ── Appendix: unmatched suggestions ──────────────────────────────────────
    if unmatched:
        doc.add_page_break()
        appendix_heading = doc.add_heading("Appendix — Suggestions Not Auto-Located", level=2)
        appendix_heading.runs[0].font.color.rgb = AMBER

        note = doc.add_paragraph()
        _add_run(
            note,
            "The following risks were identified by the AI but could not be matched "
            "to a specific clause in the extracted text with sufficient confidence. "
            "Please review and apply these manually.",
            italic=True, color=GREY, size_pt=9,
        )
        doc.add_paragraph()

        for i, risk in enumerate(unmatched, 1):
            level      = str(risk.get("risk_level", "LOW")).upper()
            clause     = risk.get("clause", "Not specified")
            why        = risk.get("why_its_a_risk", "")
            suggestion = risk.get("suggestion", "")
            level_color = RISK_COLORS.get(level, GREY)

            num_para = doc.add_paragraph()
            _add_run(num_para, f"{i}.  [{level}]  ", bold=True,
                      color=level_color, size_pt=10.5)
            _add_run(num_para, clause, italic=True, color=GREY, size_pt=10.5)

            why_para = doc.add_paragraph()
            _add_run(why_para, "Why it's a risk: ", bold=True, size_pt=10)
            _add_run(why_para, why, size_pt=10)

            sug_para = doc.add_paragraph()
            _add_run(sug_para, "Suggestion: ", bold=True, color=BLUE, size_pt=10)
            _add_run(sug_para, suggestion, color=BLUE, underline=True, size_pt=10)

            _h_rule(doc)

    # ── Serialize ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()